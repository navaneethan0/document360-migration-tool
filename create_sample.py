from docx import Document

def create_sample_docx():
    doc = Document()
    doc.add_heading('Sample Documentation Migration', 0)

    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('This is a sample document for testing the Word Document to Document360 migration tool. It contains various formatting structures.')

    doc.add_heading('Features', level=2)
    p = doc.add_paragraph('The tool supports ')
    p.add_run('bold text').bold = True
    p.add_run(', ')
    p.add_run('italic text').italic = True
    p.add_run(', and ')
    p.add_run('hyperlinks').font.color.rgb = None # Simplified representation

    doc.add_heading('Lists', level=2)
    doc.add_paragraph('Unordered List:', style='List Bullet')
    doc.add_paragraph('First item', style='List Bullet')
    doc.add_paragraph('Second item', style='List Bullet')

    doc.add_paragraph('Ordered List:', style='List Number')
    doc.add_paragraph('Step 1', style='List Number')
    doc.add_paragraph('Step 2', style='List Number')

    doc.add_heading('Data Table', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Feature'
    hdr_cells[2].text = 'Status'

    row_cells = table.rows[1].cells
    row_cells[0].text = '1'
    row_cells[1].text = 'Headings'
    row_cells[2].text = 'Supported'

    row_cells = table.rows[2].cells
    row_cells[0].text = '2'
    row_cells[1].text = 'Tables'
    row_cells[2].text = 'Supported'

    doc.save('sample.docx')
    print("sample.docx created successfully.")

if __name__ == "__main__":
    create_sample_docx()
